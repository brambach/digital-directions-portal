#!/usr/bin/env python3
"""
Generate a professional Word document for the Digital Directions Client Portal scope.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Brand color - Purple
PURPLE = RGBColor(124, 58, 237)  # #7C3AED
DARK_GRAY = RGBColor(51, 51, 51)  # #333333
LIGHT_GRAY = RGBColor(102, 102, 102)  # #666666

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc, color_hex="7C3AED"):
    """Add a horizontal line."""
    p = doc.add_paragraph()
    p_fmt = p.paragraph_format
    p_fmt.space_before = Pt(12)
    p_fmt.space_after = Pt(12)

    # Create border element
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_document():
    doc = Document()

    # Set up styles
    styles = doc.styles

    # Modify Normal style
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = DARK_GRAY

    # Create custom heading styles
    for i, size in enumerate([24, 18, 14, 12], 1):
        style_name = f'Heading {i}'
        if style_name in styles:
            style = styles[style_name]
            style.font.name = 'Calibri'
            style.font.size = Pt(size)
            style.font.color.rgb = PURPLE
            style.font.bold = True
            style.paragraph_format.space_before = Pt(18 if i == 1 else 12)
            style.paragraph_format.space_after = Pt(6)

    # ==================== TITLE PAGE ====================

    # Add some space at top
    for _ in range(4):
        doc.add_paragraph()

    # Main title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Digital Directions Client Portal")
    run.font.size = Pt(28)
    run.font.color.rgb = PURPLE
    run.font.bold = True

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Project Scope & Technical Specification")
    run.font.size = Pt(16)
    run.font.color.rgb = LIGHT_GRAY

    # Add space
    doc.add_paragraph()
    doc.add_paragraph()

    # Meta info
    meta_items = [
        ("Prepared for:", "Jack Hewitt (CEO) & Bryon (Integration Manager)"),
        ("Prepared by:", "Bryce Rambach"),
        ("Date:", "January 21, 2026"),
        ("Version:", "1.0")
    ]

    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label + " ")
        run.font.bold = True
        run.font.color.rgb = PURPLE
        run = p.add_run(value)
        run.font.color.rgb = DARK_GRAY

    # Page break after title
    doc.add_page_break()

    # ==================== TABLE OF CONTENTS ====================

    toc_heading = doc.add_heading("Table of Contents", level=1)

    toc_items = [
        ("Part 1: Executive Summary", "(For Jack)", 0),
        ("What Is This? | The Problem It Solves | Business Value | Timeline", "", 1),
        ("Part 2: Technical Specification", "(For Bryon)", 0),
        ("Architecture | Database Schema | Security | Features | Deployment", "", 1),
        ("Part 3: Delivery & Next Steps", "", 0),
    ]

    for item, note, indent in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5 * indent)
        run = p.add_run(item)
        run.font.bold = indent == 0
        if note:
            run = p.add_run(f" {note}")
            run.font.color.rgb = LIGHT_GRAY

    doc.add_page_break()

    # ==================== PART 1: EXECUTIVE SUMMARY ====================

    doc.add_heading("PART 1: EXECUTIVE SUMMARY", level=1)

    # What Is This?
    doc.add_heading("What Is This?", level=2)

    p = doc.add_paragraph()
    p.add_run("A custom-branded client portal that gives Digital Directions' clients a professional hub to:")

    features = [
        "View their project status and timeline",
        "Access project information and updates",
        "Communicate with their project team",
        "Submit and track support tickets",
        "Monitor integration health in real-time",
        "Track support hours usage"
    ]

    for feature in features:
        p = doc.add_paragraph(feature, style='List Bullet')

    p = doc.add_paragraph()
    p.add_run("Currently, DD clients juggle email threads and scattered communications. This portal consolidates everything into one clean, professional interface branded with Digital Directions' identity.")

    # The Problem It Solves
    doc.add_heading("The Problem It Solves", level=2)

    doc.add_heading("Current State (What Clients Experience)", level=3)

    current_state = [
        '"What\'s my project status?" - Emailing Jack/Tatiana for updates',
        '"When is this due?" - Unclear timelines',
        '"How do I reach my team?" - Email threads get lost',
        '"Is my integration working?" - No visibility into system health',
        '"How many support hours do I have left?" - Manual tracking'
    ]

    for item in current_state:
        doc.add_paragraph(item, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run("Result: ")
    run.font.bold = True
    p.add_run("Clients feel disorganized, DD looks less professional than competitors")

    doc.add_heading("Future State (With Portal)", level=3)

    future_state = [
        "One login - see everything",
        "Real-time project status and timeline",
        "Direct team communication",
        "Integration monitoring dashboard",
        "Support ticket system with time tracking",
        "Transparent support hours tracking"
    ]

    for item in future_state:
        doc.add_paragraph(item, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run("Result: ")
    run.font.bold = True
    p.add_run("DD looks like a premium, enterprise-grade consultancy")

    # Business Value
    doc.add_heading("Business Value", level=2)

    doc.add_heading("For Digital Directions", level=3)

    dd_benefits = [
        "Save 10+ hours per week answering repetitive status questions",
        "Improve client retention through better experience",
        "Higher perceived value - justify premium pricing",
        "Scalability - onboard new clients faster",
        "Professional brand positioning - compete with larger firms",
        "Automated support hours tracking and billing visibility"
    ]

    for item in dd_benefits:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("For DD's Clients", level=3)

    client_benefits = [
        ("Transparency", "always know project status"),
        ("Convenience", "one place for everything"),
        ("Confidence", "feel taken care of, not lost in email chains"),
        ("Accountability", "see exactly how support hours are used")
    ]

    for bold_part, rest in client_benefits:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(bold_part)
        run.font.bold = True
        p.add_run(f" - {rest}")

    # Investment & Timeline
    doc.add_heading("Investment & Timeline", level=2)

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Item"
    header_cells[1].text = "Details"
    for cell in header_cells:
        set_cell_shading(cell, "7C3AED")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True

    # Data rows
    data = [
        ("Development", "Billed at agreed hourly rate as part of ongoing work"),
        ("Timeline", "2-3 weeks from approval to initial launch"),
        ("Approach", "Incremental delivery - core features first, enhancements added based on feedback")
    ]

    for i, (item, details) in enumerate(data, 1):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = details

    doc.add_paragraph()

    # Infrastructure Costs
    doc.add_heading("Infrastructure Costs", level=2)

    p = doc.add_paragraph()
    p.add_run("All required services can start on free tiers, with no upfront subscription costs:")

    # Infrastructure costs table
    infra_table = doc.add_table(rows=5, cols=3)
    infra_table.style = 'Table Grid'

    # Header row
    infra_headers = ["Service", "Free Tier", "Notes"]
    for i, header in enumerate(infra_headers):
        cell = infra_table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "7C3AED")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True

    # Data rows
    infra_data = [
        ("Clerk (Authentication)", "10,000 monthly active users", "More than sufficient for DD's client base"),
        ("Vercel (Hosting)", "Generous bandwidth included", "Automatic scaling, no server management"),
        ("Vercel Postgres (Database)", "256 MB storage", "Plenty for portal data (no file storage)"),
        ("Resend (Email)", "3,000 emails/month", "Monitor usage; upgrade only if needed"),
    ]

    for i, (service, free_tier, notes) in enumerate(infra_data, 1):
        infra_table.rows[i].cells[0].text = service
        infra_table.rows[i].cells[1].text = free_tier
        infra_table.rows[i].cells[2].text = notes

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Bottom line: ")
    run.font.bold = True
    p.add_run("We can launch and operate the portal on free tiers initially. We'll monitor usage over the first few months and only upgrade services if/when needed. This keeps ongoing costs at $0 until the portal proves its value.")

    doc.add_paragraph()

    # What's Included
    doc.add_heading("What's Included", level=2)

    doc.add_heading("Admin Portal (DD Team View)", level=3)

    admin_features = [
        "Manage all clients from centralized dashboard",
        "Create and update projects with phase tracking",
        "Apply reusable phase templates to standardize implementations",
        "Communicate with clients via messaging system",
        "Monitor integration health across all projects",
        "Manage support tickets with time tracking",
        "Track and manage client support hour allocations",
        "Invite new users via email (invite-only access model)",
        "View project analytics and status reports"
    ]

    for item in admin_features:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Client Portal (What Clients See)", level=3)

    client_features = [
        "Clean, branded dashboard (DD purple colors and professional design)",
        "Project status and timeline visibility with visual phase stepper",
        "Direct messaging with DD team",
        "Support ticket system with response tracking",
        "Integration health monitoring (read-only view)",
        "Support hours usage visibility",
        "Mobile responsive design"
    ]

    for item in client_features:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Core Features", level=3)

    core_features = [
        "Role-based access control (admins vs clients)",
        "Invite-only user registration (no public signup)",
        "Multi-user support per client (e.g., HR Director + Payroll Manager)",
        "Real-time messaging system",
        "Project phase tracking with visual progress and reusable templates",
        "Support ticket management with priority levels",
        "Time tracking on tickets with auto-deduction from support hours",
        "Integration health monitoring for HiBob, KeyPay, NetSuite, ADP, Workato",
        "In-app notifications",
        "Mobile responsive interface"
    ]

    for item in core_features:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("What's NOT Included (Future Phases)", level=3)

    p = doc.add_paragraph()
    p.add_run("These features are excluded from initial scope but can be added in future iterations:")

    not_included = [
        "File upload/download system (security review needed)",
        "Workbook management",
        "Invoice generation (use existing QuickBooks)",
        "Advanced analytics and reporting dashboards",
        "Mobile native applications",
        "Video call integration (use Zoom/Meet)",
        "Contract signing (use DocuSign)"
    ]

    for item in not_included:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ==================== PART 2: TECHNICAL SPECIFICATION ====================

    doc.add_heading("PART 2: TECHNICAL SPECIFICATION", level=1)

    p = doc.add_paragraph()
    run = p.add_run("For Bryon - Integration Manager")
    run.font.italic = True
    run.font.color.rgb = LIGHT_GRAY

    # Architecture Overview
    doc.add_heading("Architecture Overview", level=2)
    doc.add_heading("Tech Stack", level=3)

    # Tech stack table
    table = doc.add_table(rows=10, cols=3)
    table.style = 'Table Grid'

    headers = ["Layer", "Technology", "Purpose"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "7C3AED")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True

    tech_data = [
        ("Frontend Framework", "Next.js 15 (App Router)", "Modern React framework with server-side rendering"),
        ("", "TypeScript", "Type safety and enhanced developer experience"),
        ("", "Tailwind CSS + Shadcn UI", "Professional styling and component library"),
        ("Backend", "Next.js API Routes", "Serverless API endpoints"),
        ("", "PostgreSQL (Vercel Postgres)", "Primary database"),
        ("", "Drizzle ORM", "Type-safe database queries and migrations"),
        ("Authentication", "Clerk", "Enterprise-grade auth with role-based access"),
        ("Hosting", "Vercel", "Frontend, API, and edge network hosting"),
        ("Email", "Resend", "Transactional emails for invites and notifications"),
    ]

    for i, (layer, tech, purpose) in enumerate(tech_data, 1):
        table.rows[i].cells[0].text = layer
        if layer:
            table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[i].cells[1].text = tech
        table.rows[i].cells[2].text = purpose

    doc.add_paragraph()

    # Authentication Architecture
    doc.add_heading("Authentication Architecture", level=3)

    p = doc.add_paragraph()
    run = p.add_run("Key Design Decision: ")
    run.font.bold = True
    p.add_run("Clerk is the source of truth for user profile data (email, name, avatar). The database only stores: clerkId, role, and agencyId/clientId. This separation ensures user data stays synchronized and secure.")

    auth_features = [
        "Invite-only access - No public signup; admins send email invitations",
        "Two roles: admin (DD staff) and client (client company users)",
        "Multi-user per client - Multiple portal users can access the same client's data",
        "JWT-based sessions - Secure cookie storage with automatic token refresh"
    ]

    for item in auth_features:
        p = doc.add_paragraph(style='List Bullet')
        parts = item.split(" - ")
        if len(parts) == 2:
            run = p.add_run(parts[0])
            run.font.bold = True
            p.add_run(f" - {parts[1]}")
        else:
            p.add_run(item)

    # Database Schema
    doc.add_heading("Database Schema", level=2)
    doc.add_heading("Core Tables", level=3)

    # Users table
    doc.add_heading("users", level=4)

    users_table = doc.add_table(rows=8, cols=3)
    users_table.style = 'Table Grid'

    for i, header in enumerate(["Column", "Type", "Description"]):
        cell = users_table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "7C3AED")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True

    users_data = [
        ("id", "uuid, primary key", "Internal identifier"),
        ("clerkId", "string, unique", "Reference to Clerk user (source of truth for profile)"),
        ("role", "enum: 'admin', 'client'", "Access level"),
        ("agencyId", "uuid, foreign key", "Link to agency (for admins)"),
        ("clientId", "uuid, foreign key", "Link to client company (for client users)"),
        ("deletedAt", "timestamp, nullable", "Soft delete marker"),
        ("createdAt, updatedAt", "timestamps", "Audit fields"),
    ]

    for i, (col, type_, desc) in enumerate(users_data, 1):
        users_table.rows[i].cells[0].text = col
        users_table.rows[i].cells[1].text = type_
        users_table.rows[i].cells[2].text = desc

    doc.add_paragraph()

    # Clients table
    doc.add_heading("clients", level=4)

    clients_table = doc.add_table(rows=11, cols=3)
    clients_table.style = 'Table Grid'

    for i, header in enumerate(["Column", "Type", "Description"]):
        cell = clients_table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "7C3AED")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True

    clients_data = [
        ("id", "uuid, primary key", "Identifier"),
        ("agencyId", "uuid, foreign key", "Link to Digital Directions"),
        ("companyName", "string", "Client company name"),
        ("contactName", "string", "Primary contact"),
        ("contactEmail", "string", "Primary contact email"),
        ("status", "enum: 'active', 'inactive', 'archived'", "Client status"),
        ("supportHoursPerMonth", "integer", "Allocated minutes per month"),
        ("hoursUsedThisMonth", "integer", "Minutes used in current billing period"),
        ("supportBillingCycleStart", "date", "When current billing period started"),
        ("deletedAt", "timestamp, nullable", "Soft delete marker"),
    ]

    for i, (col, type_, desc) in enumerate(clients_data, 1):
        clients_table.rows[i].cells[0].text = col
        clients_table.rows[i].cells[1].text = type_
        clients_table.rows[i].cells[2].text = desc

    doc.add_paragraph()

    # Projects table
    doc.add_heading("projects", level=4)

    projects_table = doc.add_table(rows=10, cols=3)
    projects_table.style = 'Table Grid'

    for i, header in enumerate(["Column", "Type", "Description"]):
        cell = projects_table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "7C3AED")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True

    projects_data = [
        ("id", "uuid, primary key", "Identifier"),
        ("clientId", "uuid, foreign key", "Link to client"),
        ("name", "string", "Project name"),
        ("description", "text", "Project description"),
        ("status", "enum: 'planning', 'in_progress', 'review', 'completed', 'on_hold'", "Current status"),
        ("currentPhaseId", "uuid, foreign key", "Currently active phase"),
        ("startDate", "date", "Project start"),
        ("dueDate", "date, nullable", "Target completion"),
        ("deletedAt", "timestamp, nullable", "Soft delete marker"),
    ]

    for i, (col, type_, desc) in enumerate(projects_data, 1):
        projects_table.rows[i].cells[0].text = col
        projects_table.rows[i].cells[1].text = type_
        projects_table.rows[i].cells[2].text = desc

    doc.add_paragraph()

    # Integration Monitors table
    doc.add_heading("integrationMonitors", level=4)

    p = doc.add_paragraph()
    run = p.add_run("Key Design Decision: ")
    run.font.bold = True
    p.add_run("Integration monitoring is at the project level (not client level) because integrations vary per project.")

    int_table = doc.add_table(rows=10, cols=3)
    int_table.style = 'Table Grid'

    for i, header in enumerate(["Column", "Type", "Description"]):
        cell = int_table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "7C3AED")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True

    int_data = [
        ("id", "uuid, primary key", "Identifier"),
        ("projectId", "uuid, foreign key", "Link to project (not client)"),
        ("integrationType", "string: 'hibob', 'keypay', etc.", "Integration being monitored"),
        ("name", "string", "Display name"),
        ("status", "enum: 'healthy', 'degraded', 'down', 'unknown'", "Current status"),
        ("enabled", "boolean", "Whether monitoring is active"),
        ("checkIntervalMinutes", "integer", "How often to check (default: 5)"),
        ("lastCheckedAt", "timestamp", "Last health check time"),
        ("workatoCredentials", "text, nullable", "Encrypted Workato API credentials"),
    ]

    for i, (col, type_, desc) in enumerate(int_data, 1):
        int_table.rows[i].cells[0].text = col
        int_table.rows[i].cells[1].text = type_
        int_table.rows[i].cells[2].text = desc

    doc.add_paragraph()

    # Other tables summary
    doc.add_heading("Additional Tables", level=4)

    p = doc.add_paragraph()
    p.add_run("The database also includes these tables:")

    other_tables = [
        "agencies - Digital Directions company record (single-tenant)",
        "projectPhases - Individual phases within projects",
        "phaseTemplates & templatePhases - Reusable phase definitions",
        "messages - Project-level messaging",
        "tickets - Support tickets with priority and status",
        "ticketComments - Threaded ticket responses (supports internal notes)",
        "ticketTimeEntries - Time tracking per ticket",
        "invites - Email invitations with 7-day expiry",
        "userNotifications - In-app notification system",
        "supportHourLogs - Historical support hours tracking"
    ]

    for item in other_tables:
        p = doc.add_paragraph(style='List Bullet')
        parts = item.split(" - ")
        run = p.add_run(parts[0])
        run.font.bold = True
        if len(parts) > 1:
            p.add_run(f" - {parts[1]}")

    # Security
    doc.add_heading("Security & Access Control", level=2)

    doc.add_heading("Authentication", level=3)
    auth_list = [
        "Clerk manages all session handling and token refresh",
        "JWT-based authentication for API routes",
        "Secure cookie storage (httpOnly, sameSite)",
        "Password requirements enforced by Clerk",
        "Account verification via email",
        "Invite-only registration (no public signup)"
    ]
    for item in auth_list:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Authorization", level=3)
    authz_list = [
        "Admin users: Full access to all clients, projects, and data",
        "Client users: Restricted to their assigned clientId only",
        "Enforced at both API route level and database query level",
        "All queries include soft-delete filter (deletedAt IS NULL)"
    ]
    for item in authz_list:
        p = doc.add_paragraph(style='List Bullet')
        if ":" in item:
            parts = item.split(": ")
            run = p.add_run(parts[0] + ": ")
            run.font.bold = True
            p.add_run(parts[1])
        else:
            p.add_run(item)

    doc.add_heading("Data Protection", level=3)
    protection_list = [
        "HTTPS enforced on all connections (Vercel default)",
        "Environment variables secured (never committed to repository)",
        "SQL injection prevention via parameterized queries (Drizzle ORM)",
        "XSS protection via React's built-in content escaping",
        "CSRF protection on all state-changing operations",
        "Workato credentials encrypted using AES-256-GCM"
    ]
    for item in protection_list:
        doc.add_paragraph(item, style='List Bullet')

    # Key Features
    doc.add_heading("Key Features Implementation", level=2)

    # Feature 1
    doc.add_heading("1. Project Phase Tracking", level=3)

    p = doc.add_paragraph()
    run = p.add_run("Visual Timeline: ")
    run.font.bold = True
    p.add_run("Stepper component displays current project phase with standard phases (Discovery → Build → Testing → UAT → Go Live). Reusable phase templates allow standardized implementations with custom phases per project.")

    p = doc.add_paragraph()
    run = p.add_run("Status Updates: ")
    run.font.bold = True
    p.add_run("Admin updates phase status in real-time, clients see changes immediately, notifications sent on phase transitions.")

    # Feature 2
    doc.add_heading("2. Messaging System", level=3)

    p = doc.add_paragraph()
    p.add_run("Thread-based conversations per project with 30-second polling for real-time updates. Unread message indicators, automatic notification creation, and indefinite message history.")

    # Feature 3
    doc.add_heading("3. Integration Health Monitoring", level=3)

    p = doc.add_paragraph()
    run = p.add_run("Supported Integrations: ")
    run.font.bold = True
    p.add_run("HiBob, KeyPay, NetSuite, ADP, Workato")

    p = doc.add_paragraph()
    run = p.add_run("How It Works: ")
    run.font.bold = True
    p.add_run("Vercel Cron job runs every 5 minutes, checking all enabled monitors via status page APIs. Status indicators (Healthy/Degraded/Down) with historical metrics. Configurable alert thresholds with email and in-app notifications.")

    # Feature 4
    doc.add_heading("4. Support Hours Tracking", level=3)

    p = doc.add_paragraph()
    p.add_run("Admin sets monthly allocation per client. When logging time to tickets, hours auto-deduct from client balance. Visual progress bar shows usage vs. allocation. Historical logs maintained per billing period.")

    # Feature 5
    doc.add_heading("5. Support Ticket System", level=3)

    p = doc.add_paragraph()
    run = p.add_run("Ticket Types: ")
    run.font.bold = True
    p.add_run("general_support, project_issue, feature_request, bug_report")

    p = doc.add_paragraph()
    run = p.add_run("Workflow: ")
    run.font.bold = True
    p.add_run("Client submits → Admin claims/assigns → Threaded comments → Time logging → Resolution with notes")

    p = doc.add_paragraph()
    run = p.add_run("Features: ")
    run.font.bold = True
    p.add_run("Priority-based queue, internal comments (hidden from clients), time tracking per ticket, Slack notifications")

    # Deployment
    doc.add_heading("Deployment Strategy", level=2)

    deploy_table = doc.add_table(rows=4, cols=3)
    deploy_table.style = 'Table Grid'

    for i, header in enumerate(["Environment", "Purpose", "URL Pattern"]):
        cell = deploy_table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "7C3AED")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True

    deploy_data = [
        ("Development", "Local development", "localhost:3000"),
        ("Preview", "PR review & testing", "*.vercel.app (auto-generated)"),
        ("Production", "Live client access", "portal.digitaldirections.com.au"),
    ]

    for i, (env, purpose, url) in enumerate(deploy_data, 1):
        deploy_table.rows[i].cells[0].text = env
        deploy_table.rows[i].cells[1].text = purpose
        deploy_table.rows[i].cells[2].text = url

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Automatic deployments from GitHub, preview deployments for every PR, zero-downtime production deployments with instant rollback capability.")

    doc.add_page_break()

    # ==================== PART 3: DELIVERY & NEXT STEPS ====================

    doc.add_heading("PART 3: DELIVERY & NEXT STEPS", level=1)

    # Development Approach
    doc.add_heading("Development Approach", level=2)

    p = doc.add_paragraph()
    p.add_run("This portal is being developed using Claude Code, an AI-powered development assistant that enables:")

    dev_benefits = [
        ("Faster iteration cycles", "Features built, tested, and refined in hours instead of days"),
        ("Higher code quality", "Automated best practices and consistent patterns"),
        ("Rapid bug fixes", "Issues identified and resolved quickly"),
        ("Incremental delivery", "Working features deployed as completed")
    ]

    for bold, rest in dev_benefits:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(bold)
        run.font.bold = True
        p.add_run(f" - {rest}")

    # Timeline
    doc.add_heading("Development Timeline", level=2)

    timeline_table = doc.add_table(rows=4, cols=3)
    timeline_table.style = 'Table Grid'

    for i, header in enumerate(["Week", "Focus", "Deliverables"]):
        cell = timeline_table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "7C3AED")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True

    timeline_data = [
        ("Week 1", "Foundation & Core Features", "Database setup, Clerk auth, layouts, role-based routing, project management, dashboard stats"),
        ("Week 2", "Communication & Monitoring", "Messaging, notifications, tickets, integration monitoring, support hours"),
        ("Week 3", "Polish, Testing & Launch", "UI/UX refinements, testing, bug fixes, optimization, training, deployment"),
    ]

    for i, (week, focus, deliverables) in enumerate(timeline_data, 1):
        timeline_table.rows[i].cells[0].text = week
        timeline_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        timeline_table.rows[i].cells[1].text = focus
        timeline_table.rows[i].cells[2].text = deliverables

    doc.add_paragraph()

    # Deliverables
    doc.add_heading("Deliverables", level=2)

    p = doc.add_paragraph()
    p.add_run("Upon completion, DD will receive:")

    deliverables = [
        "Fully functional portal (admin + client views)",
        "Deployed to production with live URL",
        "Database seeded with current DD clients",
        "Team training session (1 hour)",
        "Admin documentation",
        "Client onboarding guide",
        "Post-launch support (2 weeks intensive)"
    ]

    for item in deliverables:
        doc.add_paragraph(item, style='List Bullet')

    # Implementation Process
    doc.add_heading("Implementation Process", level=2)

    process_steps = [
        ("Review & Approve Scope", "Jack & Bryon review this document, provide feedback"),
        ("Kickoff Meeting", "Align on priorities, gather current client list and project data"),
        ("Development Sprint", "2-3 weeks of building with regular progress updates"),
        ("Testing & Training", "DD team tests portal, provides feedback, receives training"),
        ("Launch", "Deploy to production, begin onboarding first clients"),
        ("Support Period", "2 weeks of intensive support, then ongoing maintenance")
    ]

    for i, (step, desc) in enumerate(process_steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {step}")
        run.font.bold = True
        p.add_run(f" - {desc}")

    # Success Criteria
    doc.add_heading("Success Criteria", level=2)

    p = doc.add_paragraph()
    p.add_run("The portal will be considered successfully launched when:")

    criteria = [
        "All core features are functional and tested",
        "DD team can independently manage clients and projects",
        "At least 3 clients successfully onboarded and using portal",
        "No critical bugs or security issues identified",
        "Team training completed",
        "Documentation delivered"
    ]

    for item in criteria:
        doc.add_paragraph(item, style='List Bullet')

    # Future Roadmap
    doc.add_heading("Future Roadmap", level=2)

    doc.add_heading("Phase 2 (1-2 weeks after launch)", level=3)
    phase2 = [
        "Email notifications (in addition to in-app)",
        "Advanced project filtering and search",
        "Bulk operations for admin users",
        "Client feedback surveys",
        "Enhanced analytics dashboard"
    ]
    for item in phase2:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Phase 3 (1-2 months after launch)", level=3)
    phase3 = [
        "File upload/download system (pending security review)",
        "Advanced file organization (folders, categories)",
        "Custom reporting capabilities",
        "API for third-party integrations"
    ]
    for item in phase3:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Phase 4 (3-6 months after launch)", level=3)
    phase4 = [
        "Mobile native applications (iOS/Android)",
        "Advanced workflow automation",
        "Client self-service features",
        "Integration with additional DD tools"
    ]
    for item in phase4:
        doc.add_paragraph(item, style='List Bullet')

    # Footer
    add_horizontal_line(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Document prepared by Bryce Rambach | Digital Directions Portal v1.0 | January 2026")
    run.font.color.rgb = LIGHT_GRAY
    run.font.italic = True

    return doc

if __name__ == "__main__":
    doc = create_document()
    output_path = "/Users/bryce/Projects/Web Development/Digital Directions Portal/Digital_Directions_Portal_Scope.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
