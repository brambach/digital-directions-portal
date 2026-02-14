#!/usr/bin/env python3
"""Generate Security Architecture Summary as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], 'E5E7EB')

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = str(cell_text)

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    return table

def create_document():
    doc = Document()

    # Title
    title = doc.add_heading('Digital Directions Portal', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Security Architecture Summary')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    meta = doc.add_paragraph('Prepared for ISO 27001 Compliance Review | January 2026')
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(11)
    meta.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    doc.add_paragraph()

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'The Digital Directions Portal is a single-tenant Next.js application built with security-first '
        'architecture. Authentication is delegated to Clerk (SOC 2 Type II compliant), data is stored in '
        'Vercel Postgres with TLS encryption, and files are managed through UploadThing. The application '
        'implements role-based access control, credential encryption, and webhook signature verification.'
    )

    assessment = doc.add_paragraph()
    assessment.add_run('Overall Assessment: ').bold = True
    assessment.add_run('Strong technical controls with gaps in audit logging and rate limiting.')

    # Section 1: Architecture Overview
    doc.add_heading('1. Architecture Overview', level=1)

    doc.add_paragraph(
        'The portal follows a modern serverless architecture with clear separation of concerns:'
    )

    arch_items = [
        ('Client Browser', 'User interface via Next.js React application'),
        ('Vercel Edge', 'HTTPS/TLS termination, middleware for route protection and RBAC'),
        ('Clerk', 'Authentication provider (user accounts, sessions, MFA) - SOC 2 Type II'),
        ('Vercel Postgres', 'Primary data store (clients, projects, tickets, messages)'),
        ('UploadThing', 'File storage service (project files, agency logos)'),
    ]

    for component, description in arch_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{component}: ').bold = True
        p.add_run(description)

    # Section 2: Authentication & Access Control
    doc.add_heading('2. Authentication & Access Control', level=1)

    doc.add_heading('Authentication Flow', level=2)
    add_table(doc,
        ['Step', 'Component', 'Security Control'],
        [
            ('1', 'User login', 'Clerk handles authentication (supports MFA)'),
            ('2', 'Session created', 'Secure HTTP-only cookies, 24h expiration'),
            ('3', 'Request made', 'Middleware validates session token'),
            ('4', 'Role checked', 'Database stores authoritative role'),
            ('5', 'Route access', 'RBAC enforces admin vs client separation'),
        ],
        [0.5, 1.5, 4.0]
    )
    doc.add_paragraph()

    doc.add_heading('Role-Based Access Control', level=2)
    doc.add_paragraph('Two distinct roles with separate permissions:')

    p = doc.add_paragraph()
    p.add_run('ADMIN ROLE ').bold = True
    p.add_run('(Digital Directions staff):')
    admin_perms = ['Access to /dashboard/admin/*', 'Create/edit/delete clients', 'Create/edit projects',
                   'Manage support tickets', 'Configure integrations', 'Invite users']
    for perm in admin_perms:
        doc.add_paragraph(perm, style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('CLIENT ROLE ').bold = True
    p.add_run('(Client company users):')
    client_perms = ['Access to /dashboard/client/*', 'View own projects (read-only)',
                    'Submit support tickets', 'Send messages', 'Download files']
    for perm in client_perms:
        doc.add_paragraph(perm, style='List Bullet')

    doc.add_heading('Invite-Only Model', level=2)
    invite_points = [
        'No public signup - all users must be invited by admin',
        'Invite tokens: 256-bit cryptographically random (32 bytes)',
        '7-day expiration with single-use enforcement',
        'Role and client association set at invite time',
    ]
    for point in invite_points:
        doc.add_paragraph(point, style='List Bullet')

    # Section 3: Data Protection
    doc.add_heading('3. Data Protection', level=1)

    doc.add_heading('Data Storage Locations', level=2)
    add_table(doc,
        ['Data Type', 'Storage', 'Encryption', 'Access Control'],
        [
            ('User credentials', 'Clerk', 'Clerk-managed', 'N/A (delegated)'),
            ('User profiles', 'Clerk', 'Clerk-managed', 'Via Clerk API'),
            ('User roles', 'Vercel Postgres', 'TLS in transit', 'Database queries'),
            ('Client/project data', 'Vercel Postgres', 'TLS in transit', 'Role-based'),
            ('Workato API credentials', 'Vercel Postgres', 'AES-256-GCM', 'Admin only'),
            ('Files', 'UploadThing', 'UploadThing-managed', 'Project-based'),
        ],
        [1.8, 1.5, 1.5, 1.5]
    )
    doc.add_paragraph()

    doc.add_heading('Encryption Standards', level=2)
    encryption_points = [
        'In Transit: TLS 1.2+ (Vercel enforces HTTPS)',
        'At Rest (Credentials): AES-256-GCM with authenticated encryption',
        'At Rest (Database): Vercel Postgres default encryption',
    ]
    for point in encryption_points:
        doc.add_paragraph(point, style='List Bullet')

    doc.add_heading('Credential Encryption Details', level=2)
    add_table(doc,
        ['Property', 'Value'],
        [
            ('Algorithm', 'AES-256-GCM'),
            ('Key Size', '256 bits (32 bytes)'),
            ('IV', '16 bytes random per encryption'),
            ('Auth Tag', '16 bytes (prevents tampering)'),
            ('Format', 'iv:authTag:encryptedData (base64)'),
        ],
        [2.0, 4.0]
    )
    doc.add_paragraph()

    doc.add_heading('Data Minimization', level=2)
    minimization_points = [
        'Database stores only: clerkId, role, agencyId/clientId',
        'Profile data (name, email, avatar) fetched from Clerk on-demand',
        'No password storage (Clerk handles)',
        'No PII duplication between systems',
    ]
    for point in minimization_points:
        doc.add_paragraph(point, style='List Bullet')

    doc.add_heading('Soft Deletes', level=2)
    doc.add_paragraph(
        'All records use deletedAt timestamp to preserve audit trail. Queries filter WHERE deletedAt IS NULL. '
        'Only clients support permanent deletion (with cascade to related records).'
    )

    # Section 4: Third-Party Services
    doc.add_heading('4. Third-Party Services', level=1)

    add_table(doc,
        ['Service', 'Purpose', 'Compliance', 'Data Shared'],
        [
            ('Clerk', 'Authentication', 'SOC 2 Type II', 'User accounts, sessions'),
            ('Vercel', 'Hosting & Database', 'SOC 2 Type II', 'All application data'),
            ('UploadThing', 'File storage', 'SOC 2', 'Uploaded files'),
            ('Resend', 'Email delivery', 'SOC 2', 'Email addresses, content'),
            ('Slack', 'Notifications', 'SOC 2', 'Ticket/message summaries'),
        ],
        [1.2, 1.5, 1.3, 2.3]
    )
    doc.add_paragraph()

    doc.add_heading('Webhook Security', level=2)
    webhook_points = [
        'Clerk Webhooks: Svix signature verification (HMAC-SHA256)',
        'Cron Endpoints: Bearer token validation (CRON_SECRET)',
    ]
    for point in webhook_points:
        doc.add_paragraph(point, style='List Bullet')

    # Section 5: API Security
    doc.add_heading('5. API Security', level=1)

    doc.add_heading('Protection Layers', level=2)
    protection_layers = [
        ('HTTPS Only', 'Vercel enforces TLS, redirects HTTP'),
        ('Authentication', 'All protected routes require valid Clerk session'),
        ('Authorization', 'Role check on every API call'),
        ('Input Validation', 'Required field checks, type validation'),
        ('Parameterized Queries', 'Drizzle ORM prevents SQL injection'),
    ]
    for i, (layer, desc) in enumerate(protection_layers, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(f'{layer}: ').bold = True
        p.add_run(desc)

    doc.add_heading('Public Endpoints', level=2)
    doc.add_paragraph('Intentionally exposed endpoints with their protection mechanisms:')
    add_table(doc,
        ['Endpoint', 'Protection'],
        [
            ('/api/webhooks/clerk', 'Svix signature verification'),
            ('/api/invites/validate', 'Token validation only'),
            ('/api/invites/accept', 'Token validation + Clerk session'),
            ('/api/cron/*', 'Bearer token (CRON_SECRET)'),
        ],
        [2.5, 3.5]
    )
    doc.add_paragraph()

    doc.add_heading('Error Handling', level=2)
    error_points = [
        'Generic 500 errors returned (no stack traces exposed)',
        'Specific errors only for validation (400) and auth (401/403)',
        'Full errors logged server-side only',
    ]
    for point in error_points:
        doc.add_paragraph(point, style='List Bullet')

    # Section 6: File Upload Security
    doc.add_heading('6. File Upload Security', level=1)

    doc.add_heading('Controls', level=2)
    upload_controls = [
        ('Authentication', 'Required for all uploads'),
        ('Authorization', 'Admin-only for agency logos'),
        ('File Type', 'Restricted to images for logos'),
        ('Size Limit', '8MB maximum'),
        ('Processing', 'UploadThing handles virus scanning'),
    ]
    for control, desc in upload_controls:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{control}: ').bold = True
        p.add_run(desc)

    doc.add_heading('Access Pattern', level=2)
    access_points = [
        'Files stored in UploadThing (not local filesystem)',
        'Database stores URL reference only',
        'Access controlled at project level',
        'No direct file system access',
    ]
    for point in access_points:
        doc.add_paragraph(point, style='List Bullet')

    # Section 7: Audit & Monitoring
    doc.add_heading('7. Audit & Monitoring', level=1)

    doc.add_heading('Current Logging', level=2)
    add_table(doc,
        ['Event', 'Logged', 'Location'],
        [
            ('User login', 'Timestamp', 'clientActivity.lastLogin'),
            ('Message sent', 'Timestamp', 'clientActivity.lastMessageSent'),
            ('File downloaded', 'Timestamp', 'clientActivity.lastFileDownloaded'),
            ('Ticket created', 'Full details', 'Slack notification'),
            ('Ticket assigned', 'Full details', 'Slack + in-app notification'),
            ('Integration alert', 'Full details', 'Email + in-app notification'),
            ('API errors', 'Full error', 'Vercel Function Logs'),
        ],
        [1.5, 1.2, 3.5]
    )
    doc.add_paragraph()

    doc.add_heading('In-App Notifications', level=2)
    notification_points = [
        'Real-time bell icon with unread count',
        'Types: integration alerts, ticket updates, messages, file uploads',
        '30-second auto-refresh',
    ]
    for point in notification_points:
        doc.add_paragraph(point, style='List Bullet')

    # Section 8: Environment & Secrets
    doc.add_heading('8. Environment & Secrets', level=1)

    doc.add_heading('Required Secrets (Production)', level=2)
    add_table(doc,
        ['Secret', 'Purpose'],
        [
            ('CLERK_SECRET_KEY', 'Clerk API authentication'),
            ('CLERK_WEBHOOK_SECRET', 'Webhook signature verification'),
            ('POSTGRES_URL', 'Database connection'),
            ('UPLOADTHING_TOKEN', 'File upload service'),
            ('CRON_SECRET', 'Cron job authentication'),
            ('CREDENTIALS_ENCRYPTION_KEY', 'AES-256 key for credentials'),
        ],
        [2.5, 3.5]
    )
    doc.add_paragraph()

    doc.add_heading('Optional Integrations', level=2)
    add_table(doc,
        ['Secret', 'Purpose'],
        [
            ('RESEND_API_KEY', 'Email notifications'),
            ('SLACK_BOT_TOKEN', 'Slack notifications'),
            ('SLACK_CHANNEL_ID', 'Slack channel target'),
        ],
        [2.5, 3.5]
    )
    doc.add_paragraph()

    doc.add_heading('Secret Management', level=2)
    secret_points = [
        'Stored in Vercel environment variables',
        'Not committed to repository',
        'Separate values for dev/staging/production',
    ]
    for point in secret_points:
        doc.add_paragraph(point, style='List Bullet')

    # Section 9: Known Gaps & Recommendations
    doc.add_heading('9. Known Gaps & Recommendations', level=1)

    doc.add_heading('High Priority', level=2)
    add_table(doc,
        ['Gap', 'Risk', 'Recommendation'],
        [
            ('No rate limiting', 'Brute force attacks', 'Add @upstash/ratelimit middleware'),
            ('Limited audit logging', 'Compliance gap', 'Create comprehensive audit log table'),
            ('Invite tokens unhashed', 'Database breach exposure', 'Hash with bcrypt before storage'),
        ],
        [1.8, 1.8, 2.7]
    )
    doc.add_paragraph()

    doc.add_heading('Medium Priority', level=2)
    add_table(doc,
        ['Gap', 'Risk', 'Recommendation'],
        [
            ('No file access logging', "Can't audit downloads", 'Log file access to clientActivity'),
            ('User enumeration possible', 'Information disclosure', 'Standardize error messages'),
            ('No input schema validation', 'Injection risk', 'Add Zod validation on API routes'),
            ('No key rotation policy', 'Long-term key exposure', 'Document rotation schedule'),
        ],
        [2.0, 1.8, 2.5]
    )
    doc.add_paragraph()

    # Section 10: ISO 27001 Control Mapping
    doc.add_heading('10. ISO 27001 Control Mapping', level=1)

    doc.add_heading('Covered Controls', level=2)
    covered = [
        ('A.5 Access Control', 'Middleware RBAC, Clerk authentication'),
        ('A.5.4 Password Management', 'Delegated to Clerk (MFA supported)'),
        ('A.7.1 Cryptographic Controls', 'AES-256-GCM for credentials'),
        ('A.8.3 Network Security', 'HTTPS/TLS enforced'),
        ('A.10.1 Information Protection', 'Soft deletes, encryption'),
    ]
    for control, impl in covered:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{control}: ').bold = True
        p.add_run(impl)

    doc.add_heading('Partial Coverage', level=2)
    partial = [
        ('A.5.3 Access Rights Management', 'No approval workflow for role changes'),
        ('A.7.2 Key Management', 'No documented rotation policy'),
        ('A.11.1 Audit Logging', 'Limited to activity timestamps'),
    ]
    for control, gap in partial:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{control}: ').bold = True
        p.add_run(gap)

    doc.add_heading('Gaps Requiring Attention', level=2)
    gaps = [
        ('A.6.1 Security Roles', 'No formal ISMS roles defined'),
        ('A.6.2 Security Awareness', 'No training program'),
        ('A.9.1 Incident Reporting', 'No documented process'),
        ('A.9.2 Vulnerability Management', 'No patch policy'),
        ('A.11.2 Monitoring', 'No SIEM integration'),
    ]
    for control, gap in gaps:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{control}: ').bold = True
        p.add_run(gap)

    # Section 11: Production Readiness Checklist
    doc.add_heading('11. Production Readiness Checklist', level=1)

    doc.add_heading('Must Complete', level=2)
    must_complete = [
        'Set CRON_SECRET environment variable',
        'Set CREDENTIALS_ENCRYPTION_KEY (64 hex chars)',
        'Verify Clerk webhook secret is configured',
        'Configure email (Resend) or disable notifications',
        'Configure Slack or disable notifications',
        'Review Vercel Function Logs retention',
    ]
    for item in must_complete:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run('☐ ').font.size = Pt(12)
        p.add_run(item)

    doc.add_heading('Recommended', level=2)
    recommended = [
        'Implement rate limiting on API routes',
        'Add comprehensive audit logging',
        'Hash invite tokens',
        'Add Zod validation schemas',
        'Document incident response process',
        'Set up secret rotation schedule',
    ]
    for item in recommended:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run('☐ ').font.size = Pt(12)
        p.add_run(item)

    # Section 12: Key Files Reference
    doc.add_heading('12. Key Files Reference', level=1)

    add_table(doc,
        ['File', 'Purpose'],
        [
            ('src/middleware.ts', 'Route protection, RBAC enforcement'),
            ('src/lib/auth.ts', 'Auth helpers (getCurrentUser, requireAdmin)'),
            ('src/lib/crypto.ts', 'AES-256-GCM encryption for credentials'),
            ('src/lib/db/schema.ts', 'Database schema with soft deletes'),
            ('src/app/api/webhooks/clerk/route.ts', 'Clerk webhook with signature verification'),
            ('src/app/api/uploadthing/core.ts', 'File upload with auth checks'),
            ('src/app/api/cron/check-integrations/route.ts', 'Cron with bearer token auth'),
        ],
        [3.5, 3.0]
    )
    doc.add_paragraph()

    # Summary
    doc.add_heading('Summary', level=1)

    doc.add_paragraph(
        'The Digital Directions Portal implements strong technical security controls appropriate for a '
        'client-facing B2B application. The architecture follows security best practices including:'
    )

    best_practices = [
        'Defense in depth with multiple authentication/authorization layers',
        'Least privilege with separate admin/client roles',
        'Data minimization by delegating auth to Clerk',
        'Encryption for sensitive credentials',
    ]
    for practice in best_practices:
        doc.add_paragraph(practice, style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Primary areas for improvement before production:').bold = True

    improvements = [
        'Rate limiting on API endpoints',
        'Comprehensive audit logging for compliance',
        'Hashing invite tokens',
        'Documenting operational security procedures',
    ]
    for i, item in enumerate(improvements, 1):
        doc.add_paragraph(f'{i}. {item}')

    doc.add_paragraph(
        'The application is well-positioned for ISO 27001 compliance with additional work on '
        'organizational controls (policies, procedures, training) and enhanced technical logging.'
    )

    return doc

if __name__ == '__main__':
    doc = create_document()
    output_path = '/Users/bryce/Projects/Web Development/Digital Directions Portal/Security_Architecture_Summary.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
